"""本地知识库文档读取与切片工具。"""

from __future__ import annotations

import hashlib
import html
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Iterator


@dataclass(frozen=True)
class LoadedDocument:
    """标准化后的文档对象。"""

    source: str
    content: str
    file_type: str
    metadata: dict[str, Any] = field(default_factory=dict)


SUPPORTED_EXTENSIONS = {
    ".md",
    ".txt",
    ".rst",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
}
EXCLUDED_DIRECTORY_NAMES = {
    "00_metadata",
    ".git",
    ".idea",
    ".venv",
    "__pycache__",
    "chroma_db",
    "vector_store",
    "venv",
}
SCHEMA_VERSION = "kb_chunk_schema_v2"

# Stage5-Phase6-WP1：冻结的稳定 splitter / content-format refs（只导出，不改变切分行为）。
SPLITTER_REF = "structure-aware-splitter.v2"
CHUNK_CONTENT_FORMAT_REF = "kb-content-format.v1"


def chunk_policy_from(
    *,
    chunk_size: int,
    chunk_overlap: int,
    chunk_schema_version: str = SCHEMA_VERSION,
    splitter_ref: str = SPLITTER_REF,
    chunk_content_format_ref: str = CHUNK_CONTENT_FORMAT_REF,
) -> dict[str, object]:
    """构造 canonical chunk policy descriptor（由生产 build 使用）。

    只描述 chunk 策略，不执行切分；校验与 ``retrieval_index_provenance`` 一致。
    """
    from core.knowledge_base.retrieval_index_provenance import build_chunk_policy_descriptor

    return build_chunk_policy_descriptor(
        chunk_schema_version=chunk_schema_version,
        splitter_ref=splitter_ref,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        chunk_content_format_ref=chunk_content_format_ref,
    )


def _normalize_text(text: str) -> str:
    """统一换行和尾部空白，同时保留代码与配置的行首缩进。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_html_file(path: Path) -> str:
    raw = _read_text_file(path)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(
            ["script", "style", "nav", "footer", "header", "aside", "noscript"]
        ):
            tag.decompose()
        return soup.get_text("\n")
    except Exception:
        raw = re.sub(
            r"(?is)<(script|style|nav|footer|header|aside).*?>.*?</\1>",
            "",
            raw,
        )
        return html.unescape(re.sub(r"(?s)<[^>]+>", "\n", raw))


def _read_docx_file(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.rstrip()
        if not text.strip():
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        heading = re.match(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
        if heading:
            level = min(6, max(1, int(heading.group(1))))
            parts.append(f"{'#' * level} {text.strip()}")
        else:
            parts.append(text)
    return "\n\n".join(parts)


def _calculate_file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source_file:
        while block := source_file.read(1024 * 1024):
            digest.update(block)
    return digest.hexdigest()


def _detect_language(text: str) -> str:
    sample = text[:5000]
    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", sample))
    english_count = len(re.findall(r"[A-Za-z]", sample))
    if chinese_count and english_count:
        return "mixed"
    if chinese_count:
        return "zh"
    if english_count:
        return "en"
    return "unknown"


def iter_supported_files(
    base_dir: str, *, max_files: int | None = None
) -> Iterator[Path]:
    """遍历支持入库的文件并排除本地制品目录。"""
    root = Path(base_dir).resolve()
    if not root.exists():
        return
    emitted = 0
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        relative = path.relative_to(root)
        if set(relative.parts[:-1]) & EXCLUDED_DIRECTORY_NAMES:
            continue
        yield path
        emitted += 1
        if max_files is not None and emitted >= max_files:
            break


def _common_metadata(path: Path, root: Path) -> dict[str, Any]:
    source = path.relative_to(root).as_posix()
    return {
        "doc_id": hashlib.sha1(source.encode("utf-8")).hexdigest(),
        "file_name": path.name,
        "file_hash": _calculate_file_hash(path),
        "modified_time": datetime.fromtimestamp(
            path.stat().st_mtime, timezone.utc
        ).isoformat(),
    }


def _build_table_summary(
    frame: Any, *, file_name: str, sheet_name: str
) -> tuple[str, dict[str, Any]]:
    """把表格转换为字段级摘要，避免将整张大表写入向量库。"""
    row_count = int(len(frame))
    column_count = int(len(frame.columns))
    lines = [
        f"文件：{file_name}",
        f"工作表：{sheet_name}",
        f"总行数：{row_count}",
        f"总列数：{column_count}",
        "",
        "字段说明：",
    ]
    column_names = []
    for original_column in list(frame.columns)[:100]:
        column_name = str(original_column)
        column_names.append(column_name)
        series = frame[original_column]
        samples = (
            series.dropna().astype(str).head(200).drop_duplicates().head(5).tolist()
        )
        clean_samples = [
            value.replace("\r", " ").replace("\n", " ")[:80] for value in samples
        ]
        line = (
            f"- {column_name}：类型={series.dtype}，"
            f"缺失值={int(series.isna().sum())}"
        )
        if clean_samples:
            line += f"，示例值={', '.join(clean_samples)}"
        lines.append(line)
    if column_count > 100:
        lines.append(f"- 其余 {column_count - 100} 个字段未在摘要中展开。")
    return "\n".join(lines), {
        "sheet_name": sheet_name,
        "row_count": row_count,
        "column_count": column_count,
        "columns_csv": ",".join(column_names),
        "block_type": "table_summary",
        "parser_name": "pandas",
        "parser_version": "v1",
    }


def _read_table_documents(
    path: Path,
    root: Path,
    common_metadata: dict[str, Any],
) -> list[LoadedDocument]:
    import pandas as pd

    if path.suffix.lower() == ".csv":
        frame = None
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                frame = pd.read_csv(
                    path, encoding=encoding, low_memory=False, on_bad_lines="skip"
                )
                break
            except UnicodeDecodeError:
                continue
        if frame is None:
            raise ValueError(f"无法读取 CSV 文件：{path}")
        sheets = {"csv": frame}
    else:
        sheets = pd.read_excel(path, sheet_name=None)

    source = path.relative_to(root).as_posix()
    documents = []
    for sheet_name, frame in sheets.items():
        summary, table_metadata = _build_table_summary(
            frame,
            file_name=path.name,
            sheet_name=str(sheet_name),
        )
        documents.append(
            LoadedDocument(
                source=source,
                content=_normalize_text(summary),
                file_type=path.suffix.lower().lstrip("."),
                metadata={
                    **common_metadata,
                    **table_metadata,
                    "document_title": path.stem,
                    "lang": _detect_language(summary),
                },
            )
        )
    return documents


def load_document_file(path: Path, base_dir: str | Path) -> list[LoadedDocument]:
    """解析单个文件；PDF 按页、Excel 按 Sheet 返回解析单元。"""
    root = Path(base_dir).resolve()
    path = Path(path).resolve()
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return []
    source = path.relative_to(root).as_posix()
    common_metadata = _common_metadata(path, root)

    if suffix in {".csv", ".xls", ".xlsx"}:
        return _read_table_documents(path, root, common_metadata)
    if suffix == ".pdf":
        from pypdf import PdfReader

        documents = []
        for page_number, page in enumerate(PdfReader(str(path)).pages, start=1):
            content = _normalize_text(page.extract_text() or "")
            if content:
                documents.append(
                    LoadedDocument(
                        source=source,
                        content=content,
                        file_type="pdf",
                        metadata={
                            **common_metadata,
                            "document_title": path.stem,
                            "page_start": page_number,
                            "page_end": page_number,
                            "block_type": "page",
                            "lang": _detect_language(content),
                            "parser_name": "pypdf",
                            "parser_version": "v1",
                        },
                    )
                )
        return documents

    if suffix == ".docx":
        content = _normalize_text(_read_docx_file(path))
        parser_name = "python-docx"
        content_format = "markdown"
    elif suffix in {".html", ".htm"}:
        content = _normalize_text(_read_html_file(path))
        parser_name = "beautifulsoup"
        content_format = "plain"
    else:
        content = _normalize_text(_read_text_file(path))
        parser_name = "native"
        content_format = "markdown" if suffix == ".md" else "plain"
    if not content:
        return []
    title_match = re.search(r"(?m)^#\s+(.+?)\s*$", content)
    metadata = {
        **common_metadata,
        "document_title": title_match.group(1).strip() if title_match else path.stem,
        "block_type": "document",
        "lang": _detect_language(content),
        "parser_name": parser_name,
        "parser_version": "v1",
        "content_format": content_format,
    }
    return [LoadedDocument(source, content, suffix.lstrip("."), metadata)]


def load_documents(
    base_dir: str, *, max_files: int | None = None
) -> list[LoadedDocument]:
    """兼容旧调用：批量加载目录，并隔离单个文件的解析失败。"""
    documents = []
    for path in iter_supported_files(base_dir, max_files=max_files):
        try:
            documents.extend(load_document_file(path, base_dir))
        except Exception:
            continue
    return documents


def _build_markdown_sections(content: str) -> list[dict[str, Any]]:
    """按 Markdown H1～H6 构造保留标题层级的章节。"""
    matches = list(re.finditer(r"(?m)^(#{1,6})[ \t]+(.+?)[ \t]*$", content))
    if not matches:
        return [{"start": 0, "content": content, "section_path": None}]

    sections = []
    hierarchy: list[str | None] = [None] * 6
    preamble = content[: matches[0].start()].strip()
    if preamble:
        sections.append({"start": 0, "content": preamble, "section_path": None})
    for index, match in enumerate(matches):
        level = len(match.group(1))
        hierarchy[level - 1] = match.group(2).strip()
        hierarchy[level:] = [None] * (6 - level)
        end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
        sections.append(
            {
                "start": match.start(),
                "content": content[match.start() : end].strip(),
                "section_h1": hierarchy[0],
                "section_h2": hierarchy[1],
                "section_h3": hierarchy[2],
                "section_path": " > ".join(item for item in hierarchy if item) or None,
            }
        )
    return sections


def _split_text_windows(
    text: str,
    *,
    chunk_size: int,
    chunk_overlap: int,
) -> Iterator[tuple[int, int, str]]:
    """优先在段落或行边界切片，超长段落才使用字符窗口。"""
    start = 0
    while start < len(text):
        hard_end = min(len(text), start + chunk_size)
        end = hard_end
        if hard_end < len(text):
            minimum_boundary = start + max(1, chunk_size // 2)
            candidates = [
                text.rfind("\n\n", minimum_boundary, hard_end),
                text.rfind("\n", minimum_boundary, hard_end),
                text.rfind("。", minimum_boundary, hard_end),
                text.rfind(". ", minimum_boundary, hard_end),
            ]
            boundary = max(candidates)
            if boundary >= minimum_boundary:
                end = boundary + 1
        if end <= start:
            end = hard_end
        snippet = text[start:end].strip()
        if snippet:
            yield start, end, snippet
        if end >= len(text):
            break
        start = max(start + 1, end - chunk_overlap)


def _build_embedding_content(
    document: LoadedDocument,
    snippet: str,
    section_path: str | None,
) -> str:
    prefix = []
    if document.metadata.get("document_title"):
        prefix.append(f"文档：{document.metadata['document_title']}")
    if section_path:
        prefix.append(f"章节：{section_path}")
    if document.metadata.get("page_start") is not None:
        prefix.append(f"页码：{document.metadata['page_start']}")
    if document.metadata.get("sheet_name"):
        prefix.append(f"工作表：{document.metadata['sheet_name']}")
    return "\n".join(prefix) + "\n\n" + snippet if prefix else snippet


def _resolve_source_type(document: LoadedDocument) -> str:
    if document.file_type == "md" and document.source.endswith(".pdf.md"):
        return "pdf_md"
    return document.file_type


def split_documents(
    documents: Iterable[LoadedDocument],
    *,
    chunk_size: int = 1400,
    chunk_overlap: int = 180,
    ingest_batch_id: str | None = None,
) -> list[dict]:
    """将文档切分为可向量化 chunk。"""
    if chunk_size <= 0:
        raise ValueError("chunk_size 必须大于 0。")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap 不能小于 0。")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap 必须小于 chunk_size。")

    chunks: list[dict] = []
    batch_id = ingest_batch_id or datetime.now(timezone.utc).isoformat()
    source_chunk_counters: dict[str, int] = {}

    for document in documents:
        source_type = _resolve_source_type(document)
        is_markdown = (
            document.file_type == "md"
            or document.metadata.get("content_format") == "markdown"
        )
        sections = (
            _build_markdown_sections(document.content)
            if is_markdown
            else [{"start": 0, "content": document.content, "section_path": None}]
        )
        source_chunk_counters.setdefault(document.source, 0)
        for section in sections:
            embedding_prefix = _build_embedding_content(
                document,
                "",
                section.get("section_path"),
            )
            content_chunk_size = max(1, chunk_size - len(embedding_prefix))
            content_overlap = min(chunk_overlap, max(0, content_chunk_size - 1))
            for local_start, local_end, snippet in _split_text_windows(
                section["content"],
                chunk_size=content_chunk_size,
                chunk_overlap=content_overlap,
            ):
                char_start = int(section["start"]) + local_start
                char_end = int(section["start"]) + local_end
                content_hash = hashlib.sha1(snippet.encode("utf-8")).hexdigest()
                stable_key = "|".join(
                    [
                        document.source,
                        str(document.metadata.get("page_start") or ""),
                        str(document.metadata.get("sheet_name") or ""),
                        str(char_start),
                        str(char_end),
                        str(section.get("section_path") or ""),
                        content_hash,
                    ]
                )
                metadata = {
                    **document.metadata,
                    "schema_version": SCHEMA_VERSION,
                    "source": document.source,
                    "source_type": source_type,
                    "file_type": document.file_type,
                    "chunk_id": hashlib.sha1(stable_key.encode("utf-8")).hexdigest(),
                    "chunk_index": source_chunk_counters[document.source],
                    "content_hash": content_hash,
                    "char_start": char_start,
                    "char_end": char_end,
                    "offset": char_start,
                    "section_h1": section.get("section_h1"),
                    "section_h2": section.get("section_h2"),
                    "section_h3": section.get("section_h3"),
                    "section_path": section.get("section_path"),
                    "block_type": document.metadata.get("block_type", "paragraph"),
                    "lang": document.metadata.get("lang", _detect_language(snippet)),
                    "chunker_name": "structure_aware_splitter",
                    "chunker_version": "v2",
                    "ingest_batch_id": batch_id,
                }
                source_chunk_counters[document.source] += 1
                chunks.append(
                    {
                        "page_content": _build_embedding_content(
                            document,
                            snippet,
                            section.get("section_path"),
                        ),
                        "metadata": metadata,
                    }
                )
    return chunks


def ensure_test_documents(base_dir: str) -> list[str]:
    """写入一组可直接用于测试的知识库文档。"""
    root = Path(base_dir)
    root.mkdir(parents=True, exist_ok=True)
    created: list[str] = []

    md_path = root / "product_faq.md"
    md_path.write_text(
        """# 本地知识库产品 FAQ

## 计费策略
- 标准版按账号计费，月付 99 元。
- 企业版支持私有部署，按年签约。

## 数据安全
- 默认只使用本地向量数据库，不上传原文。
- 可配置审计日志保存 180 天。
""",
        encoding="utf-8",
    )
    created.append(str(md_path))

    txt_path = root / "ops_runbook.txt"
    txt_path.write_text(
        """故障排查步骤：
1. 检查 embedding 模型目录是否存在。
2. 检查 chroma_db 是否可写。
3. 如命中率低，增大 top_k 并重建索引。""",
        encoding="utf-8",
    )
    created.append(str(txt_path))

    try:
        import pandas as pd

        xlsx_path = root / "sla_matrix.xlsx"
        pd.DataFrame(
            [
                {"服务等级": "P1", "响应时效": "15分钟", "恢复目标": "2小时"},
                {"服务等级": "P2", "响应时效": "1小时", "恢复目标": "8小时"},
                {"服务等级": "P3", "响应时效": "4小时", "恢复目标": "24小时"},
            ]
        ).to_excel(xlsx_path, index=False)
        created.append(str(xlsx_path))
    except Exception:
        csv_path = root / "sla_matrix.csv"
        csv_path.write_text(
            "服务等级,响应时效,恢复目标\nP1,15分钟,2小时\nP2,1小时,8小时\nP3,4小时,24小时\n",
            encoding="utf-8",
        )
        created.append(str(csv_path))

    return created
